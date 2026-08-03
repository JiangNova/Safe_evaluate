"""Anonymous, stateless API routes for the leadership writing workbench."""

from __future__ import annotations

import json
import os
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import ValidationError

from .auth import authenticate_leadership_user, verify_token
from .config import PUBLIC_JOB_MAX_FILES, PUBLIC_JOB_MAX_TOTAL_SIZE
from .leadership_writer import (
    GeneratedDocument,
    LeadershipProfile,
    LeadershipWriterError,
    WritingTask,
    generate_document,
    revise_document,
)
from .models import (
    LeadershipDocxExportRequest,
    LeadershipProfileRequest,
    LeadershipRevisionRequest,
    LoginRequest,
    LoginResponse,
)
from .public_files import (
    SourceExtractionError,
    UploadValidationError,
    extract_source,
    validate_upload,
)


DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
router = APIRouter(tags=["ai-writing"])
leadership_security = HTTPBearer(auto_error=False)


async def require_leadership_auth(
    credentials: HTTPAuthorizationCredentials | None = Depends(leadership_security),
) -> dict:
    """Accept only tokens issued for the leadership writing workbench."""
    if credentials is None:
        raise HTTPException(status_code=401, detail="请先登录领导文稿助手")
    payload = verify_token(credentials.credentials)
    if payload is None or payload.get("role") != "leader_assistant":
        raise HTTPException(status_code=401, detail="登录已失效，请重新登录")
    return payload


@router.post("/auth/login", response_model=LoginResponse)
async def login_leadership_user(req: LoginRequest):
    """Authenticate one of the small, dedicated leadership-writing accounts."""
    token = authenticate_leadership_user(req.username, req.password)
    if token is None:
        raise HTTPException(status_code=401, detail="用户名或密码错误")
    return LoginResponse(token=token, user={"username": req.username, "role": "leader_assistant"})


def _api_error(status_code: int, code: str, message: str, *, stage: str) -> HTTPException:
    """Return the small, stable error envelope consumed by the standalone app."""
    return HTTPException(
        status_code=status_code,
        detail={"code": code, "message": message, "stage": stage},
    )


def _to_profile(raw_profile: str) -> LeadershipProfile:
    try:
        profile_data = json.loads(raw_profile)
        request_profile = LeadershipProfileRequest.model_validate(profile_data)
        return LeadershipProfile.model_validate(request_profile.model_dump())
    except (json.JSONDecodeError, ValidationError, TypeError) as exc:
        raise _api_error(
            400,
            "invalid_request",
            "身份档案格式不正确，请补全身份名称后重试。",
            stage="validation",
        ) from exc


def _to_task(task_type: str, requirement: str) -> WritingTask:
    try:
        return WritingTask(task_type=task_type, requirement=requirement)
    except ValidationError as exc:
        raise _api_error(
            400,
            "invalid_request",
            "写作任务类型或要求不正确。",
            stage="validation",
        ) from exc


async def _extract_uploaded_sources(files: list[UploadFile]) -> list:
    """Validate and parse uploads inside a request-scoped temporary directory.

    ``store_upload`` intentionally is not used here: it records files in the public
    job database for later retrieval, which would violate this workbench's no-server-
    persistence requirement.  The same validation and extraction primitives are used,
    and the temporary directory is removed before this function returns.
    """
    if len(files) > PUBLIC_JOB_MAX_FILES:
        raise UploadValidationError("本次任务的参考文件数量超过限制")

    total_size = 0
    sources = []
    with tempfile.TemporaryDirectory(prefix="leadership-assistant-") as directory:
        for index, uploaded_file in enumerate(files, start=1):
            data = await uploaded_file.read()
            total_size += len(data)
            if total_size > PUBLIC_JOB_MAX_TOTAL_SIZE:
                raise UploadValidationError("本次任务的参考文件总大小超过限制")
            validated = validate_upload(
                "basis",
                uploaded_file.filename or "",
                uploaded_file.content_type or "",
                data,
            )
            temporary_path = os.path.join(directory, f"source-{index}{validated.extension}")
            Path(temporary_path).write_bytes(validated.data)
            sources.append(
                extract_source(
                    {
                        "id": index,
                        "safe_name": os.path.basename(temporary_path),
                        "storage_path": temporary_path,
                        "original_name": validated.original_name,
                        "mime_type": validated.mime_type,
                    }
                )
            )
    return sources


@router.post("/generate", response_model=GeneratedDocument)
async def generate_leadership_document(
    profile: str = Form(...),
    task_type: str = Form(...),
    requirement: str = Form(...),
    files: list[UploadFile] | None = File(default=None),
    _auth: dict = Depends(require_leadership_auth),
):
    """Generate one Markdown draft and discard all uploaded reference material."""
    parsed_profile = _to_profile(profile)
    task = _to_task(task_type, requirement)
    try:
        sources = await _extract_uploaded_sources(files or [])
    except UploadValidationError as exc:
        raise _api_error(400, "invalid_upload", str(exc), stage="upload") from exc
    except SourceExtractionError as exc:
        raise _api_error(400, "source_extraction_failed", str(exc), stage="extraction") from exc

    try:
        return await generate_document(parsed_profile, task, sources)
    except (LeadershipWriterError, RuntimeError) as exc:
        raise _api_error(
            502,
            "generation_failed",
            "文稿生成服务暂时不可用，请稍后重试。",
            stage="generation",
        ) from exc


@router.post("/revise", response_model=GeneratedDocument)
async def revise_leadership_document(
    body: LeadershipRevisionRequest,
    _auth: dict = Depends(require_leadership_auth),
):
    """Revise a browser-supplied draft without retaining it after the response."""
    profile = LeadershipProfile.model_validate(body.profile.model_dump())
    task = WritingTask(task_type=body.task_type, requirement=body.requirement)
    existing_document = GeneratedDocument(
        title=body.title,
        content_markdown=body.content_markdown,
        warnings=body.warnings,
    )
    try:
        return await revise_document(
            profile,
            task,
            existing_document,
            body.revision_instruction,
        )
    except (LeadershipWriterError, RuntimeError) as exc:
        raise _api_error(
            502,
            "revision_failed",
            "文稿改写服务暂时不可用，请稍后重试。",
            stage="revision",
        ) from exc


def _set_run_font(run, font_name: str, font_size: float, *, bold: bool = False) -> None:
    """Set Latin and East Asian fonts explicitly so Chinese text renders predictably."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attribute), font_name)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_inline_markdown(paragraph, text: str, *, font_name: str = "宋体", font_size: float = 12, bold: bool = False) -> None:
    """Write plain text and ``**bold**`` spans without exposing Markdown markers."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if is_bold else part.replace("**", ""))
        _set_run_font(run, font_name, font_size, bold=bold or is_bold)


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    _add_inline_markdown(
        paragraph,
        text,
        font_name="黑体",
        font_size=16 if level == 1 else 14,
        bold=True,
    )


def _add_body_paragraph(document: Document, text: str, *, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if style is None:
        paragraph.paragraph_format.first_line_indent = Pt(24)
    _add_inline_markdown(paragraph, text)


def _render_markdown(document: Document, content_markdown: str, title: str) -> None:
    """Render assistant Markdown into a formal Chinese Word document."""
    _configure_document(document)
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(18)
    _add_inline_markdown(title_paragraph, title, font_name="黑体", font_size=22, bold=True)
    rendered_title = False
    for raw_line in content_markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*#*$", line)
        if heading:
            heading_text = heading.group(2).strip()
            if not rendered_title and heading_text == title:
                rendered_title = True
                continue
            _add_heading(document, heading_text, len(heading.group(1)))
            continue
        if line.startswith(("- ", "* ")):
            _add_body_paragraph(document, line[2:].strip(), style="List Bullet")
            continue
        numbered = re.match(r"^\d+[.\u3001]\s+(.+)$", line)
        if numbered:
            _add_body_paragraph(document, numbered.group(1), style="List Number")
            continue
        chinese_heading = re.match(r"^(?:[一二三四五六七八九十]+、|（[一二三四五六七八九十]+）).+", line)
        if chinese_heading:
            _add_heading(document, line, 1)
            continue
        _add_body_paragraph(document, line)


def _safe_filename(title: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\x00-\x1f]+', "_", title).strip(" ._")
    return (cleaned or "领导文稿")[:100]


def _remove_file(path: str) -> None:
    try:
        os.remove(path)
    except FileNotFoundError:
        pass


@router.post("/export/docx")
async def export_leadership_docx(
    body: LeadershipDocxExportRequest,
    background_tasks: BackgroundTasks,
    _auth: dict = Depends(require_leadership_auth),
):
    """Render browser-local Markdown as a one-off DOCX download."""
    document = Document()
    _render_markdown(document, body.content_markdown, body.title)
    with tempfile.NamedTemporaryFile(prefix="leadership-export-", suffix=".docx", delete=False) as output:
        temporary_path = output.name
    try:
        document.save(temporary_path)
    except Exception:
        _remove_file(temporary_path)
        raise
    background_tasks.add_task(_remove_file, temporary_path)
    return FileResponse(
        temporary_path,
        media_type=DOCX_MEDIA_TYPE,
        filename=f"{_safe_filename(body.title)}.docx",
        background=background_tasks,
    )
