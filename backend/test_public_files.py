"""Tests for validation, storage, and extraction of public job uploads."""

import io
import os
import tempfile
import unittest
from unittest.mock import patch

from docx import Document
from PIL import Image

from backend import public_files, public_jobs


DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Policy", level=1)
    doc.add_paragraph("The score must be at least 80.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Criterion"
    table.cell(0, 1).text = "Threshold"
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def make_image_bytes(image_format: str, size: tuple[int, int]) -> bytes:
    image = Image.new("RGB", size, color=(24, 96, 144))
    stream = io.BytesIO()
    image.save(stream, format=image_format)
    return stream.getvalue()


class PublicFileTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.db_path = os.path.join(self.temp_dir.name, "jobs.db")
        self.storage_path = os.path.join(self.temp_dir.name, "storage")
        self.db_patch = patch.object(public_jobs, "DB_PATH", self.db_path)
        self.storage_patch = patch.object(
            public_files, "PUBLIC_JOB_STORAGE_DIR", self.storage_path
        )
        self.db_patch.start()
        self.storage_patch.start()
        public_jobs.init_public_job_db()
        self.job, _ = public_jobs.create_job("Assess material")

    def tearDown(self):
        self.storage_patch.stop()
        self.db_patch.stop()
        self.temp_dir.cleanup()

    def test_kind_specific_allowed_uploads(self):
        cases = [
            ("material", "photo.png", "image/png", b"\x89PNG\r\n\x1a\nrest"),
            ("material", "brief.docx", DOCX_MIME, make_docx_bytes()),
            ("basis", "policy.txt", "text/plain", "要求达到80分".encode()),
            ("template", "output.pdf", "application/pdf", b"%PDF-1.7\n%%EOF"),
        ]

        for kind, name, mime, data in cases:
            with self.subTest(kind=kind, name=name):
                validated = public_files.validate_upload(kind, name, mime, data)
                self.assertEqual(validated.extension, os.path.splitext(name)[1])

    def test_disallowed_or_mismatched_files_are_rejected(self):
        with self.assertRaises(public_files.UploadValidationError):
            public_files.validate_upload(
                "template", "macro.docm", DOCX_MIME, b"PK\x03\x04"
            )
        with self.assertRaises(public_files.UploadValidationError):
            public_files.validate_upload(
                "template", "fake.pdf", "application/pdf", b"MZ executable"
            )
        with self.assertRaises(public_files.UploadValidationError):
            public_files.validate_upload(
                "basis", "encrypted.pdf", "application/pdf", b"%PDF /Encrypt"
            )

    def test_store_upload_generates_safe_name_and_metadata(self):
        upload = public_files.validate_upload(
            "basis", "../policy.txt", "text/plain", b"Policy text"
        )
        record = public_files.store_upload(self.job["id"], upload)

        self.assertNotIn("..", record["safe_name"])
        self.assertTrue(os.path.isfile(record["storage_path"]))
        self.assertEqual(record["original_name"], "policy.txt")
        self.assertTrue(
            os.path.abspath(record["storage_path"]).startswith(
                os.path.abspath(self.storage_path)
            )
        )

    def test_docx_and_text_sources_have_stable_references(self):
        docx_record = public_files.store_upload(
            self.job["id"],
            public_files.validate_upload(
                "basis", "policy.docx", DOCX_MIME, make_docx_bytes()
            ),
        )
        text_record = public_files.store_upload(
            self.job["id"],
            public_files.validate_upload(
                "basis", "notes.txt", "text/plain", "第一条要求".encode("utf-8")
            ),
        )

        docx_source = public_files.extract_source(docx_record)
        text_source = public_files.extract_source(text_record)

        self.assertTrue(docx_source.chunks[0].source_ref.startswith("paragraph:"))
        self.assertIn("Criterion", "\n".join(c.text for c in docx_source.chunks))
        self.assertEqual(text_source.chunks[0].source_ref, "text:1")

    def test_pdf_source_has_page_references(self):
        record = public_files.store_upload(
            self.job["id"],
            public_files.validate_upload(
                "basis", "policy.pdf", "application/pdf", b"%PDF-1.7\n%%EOF"
            ),
        )
        with patch.object(
            public_files,
            "_read_pdf_pages",
            return_value=["Page one", "Page two"],
        ):
            source = public_files.extract_source(record)

        self.assertEqual(
            [chunk.source_ref for chunk in source.chunks], ["page:1", "page:2"]
        )

    def test_image_material_is_returned_as_visual_source(self):
        record = public_files.store_upload(
            self.job["id"],
            public_files.validate_upload(
                "material", "photo.png", "image/png", b"\x89PNG\r\n\x1a\nrest"
            ),
        )
        source = public_files.extract_source(record)
        self.assertEqual(source.media_type, "image")
        self.assertEqual(source.chunks[0].source_ref, "image:photo.png")

    def test_large_jpeg_uses_a_smaller_evaluation_copy_without_changing_original(self):
        original = make_image_bytes("JPEG", (3200, 1600))
        path = os.path.join(self.temp_dir.name, "large.jpg")
        with open(path, "wb") as output:
            output.write(original)
        record = {
            "kind": "material",
            "safe_name": "large.jpg",
            "mime_type": "image/jpeg",
            "storage_path": path,
        }

        prepared, mime_type = public_files.prepare_evaluation_image(record)

        self.assertEqual(mime_type, "image/jpeg")
        with Image.open(io.BytesIO(prepared)) as image:
            self.assertEqual(image.size, (2560, 1280))
        with open(path, "rb") as original_file:
            self.assertEqual(original_file.read(), original)

    def test_small_jpeg_and_png_use_the_original_bytes(self):
        for name, mime_type, image_format in (
            ("small.jpg", "image/jpeg", "JPEG"),
            ("diagram.png", "image/png", "PNG"),
        ):
            with self.subTest(name=name):
                original = make_image_bytes(image_format, (800, 600))
                path = os.path.join(self.temp_dir.name, name)
                with open(path, "wb") as output:
                    output.write(original)

                prepared, prepared_mime = public_files.prepare_evaluation_image(
                    {
                        "kind": "material",
                        "safe_name": name,
                        "mime_type": mime_type,
                        "storage_path": path,
                    }
                )

                self.assertEqual(prepared, original)
                self.assertEqual(prepared_mime, mime_type)

    def test_invalid_jpeg_falls_back_to_the_original_bytes(self):
        original = b"\xff\xd8\xffnot-a-real-jpeg"
        path = os.path.join(self.temp_dir.name, "broken.jpg")
        with open(path, "wb") as output:
            output.write(original)

        prepared, mime_type = public_files.prepare_evaluation_image(
            {
                "kind": "material",
                "safe_name": "broken.jpg",
                "mime_type": "image/jpeg",
                "storage_path": path,
            }
        )

        self.assertEqual(prepared, original)
        self.assertEqual(mime_type, "image/jpeg")


if __name__ == "__main__":
    unittest.main()
