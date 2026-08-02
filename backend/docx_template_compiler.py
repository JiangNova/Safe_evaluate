"""Deterministic structural compiler for business-form DOCX templates."""

from __future__ import annotations

import hashlib
import re
from collections.abc import Callable
from pathlib import Path

from docx import Document
from pydantic import BaseModel, Field

from .template_ir import CompiledField, CompiledTemplate, Placement


class PlacementCandidate(BaseModel):
    kind: str
    label: str
    part: str = "document"
    paragraph_index: int | None = None
    run_start: int | None = None
    run_end: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    placeholder: str | None = None
    options: list[str] = Field(default_factory=list)
    context: str = ""
    fingerprint: str


def _normalized(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _fingerprint(part: str, path: str, context: str) -> str:
    return hashlib.sha256(f"{part}|{path}|{_normalized(context)}".encode("utf-8")).hexdigest()


def _label_before(text: str, start: int) -> str:
    prefix = text[:start].rstrip()
    match = re.search(r"([^：:，,。；;\s]{1,24})[：:]?$", prefix)
    return match.group(1) if match else "待填写内容"


def _paragraph_candidates(paragraph, index: int, part: str) -> list[PlacementCandidate]:
    text = paragraph.text
    items: list[PlacementCandidate] = []
    for run_index, run in enumerate(paragraph.runs):
        if run.font.underline and (not run.text.strip() or re.fullmatch(r"[_＿\s]{2,}", run.text or "")):
            label = _label_before(text, max(0, text.find(run.text)))
            items.append(PlacementCandidate(
                kind="run_range_replace", label=label, part=part,
                paragraph_index=index, run_start=run_index, run_end=run_index,
                context=text, fingerprint=_fingerprint(part, f"p{index}:r{run_index}", text),
            ))
    for match in re.finditer(r"\{\{\s*([A-Za-z][A-Za-z0-9_]*)\s*\}\}", text):
        items.append(PlacementCandidate(
            kind="placeholder_replace", label=match.group(1), placeholder=match.group(0),
            part=part, paragraph_index=index, context=text,
            fingerprint=_fingerprint(part, f"p{index}:placeholder:{match.group(1)}", text),
        ))
    if re.search(r"年\s*[_＿\s]{2,}\s*月\s*[_＿\s]{2,}\s*日", text):
        items.append(PlacementCandidate(
            kind="date_parts", label=_label_before(text, text.find("年")), part=part,
            paragraph_index=index, context=text,
            fingerprint=_fingerprint(part, f"p{index}:date", text),
        ))
    marks = list(re.finditer(r"[□☐☑■]([^□☐☑■，,。；;\n]{1,24})", text))
    if len(marks) >= 2:
        options = [_normalized(match.group(1)) for match in marks]
        items.append(PlacementCandidate(
            kind="checkbox_select", label=_label_before(text, marks[0].start()), options=options,
            part=part, paragraph_index=index, context=text,
            fingerprint=_fingerprint(part, f"p{index}:checkbox", text),
        ))
    if not items:
        for match in re.finditer(r"(?<=[:：])\s{3,}", text):
            items.append(PlacementCandidate(
                kind="run_range_replace", label=_label_before(text, match.start()), part=part,
                paragraph_index=index, context=text,
                fingerprint=_fingerprint(part, f"p{index}:blank:{match.start()}", text),
            ))
    return items


def extract_docx_candidates(path: str) -> list[PlacementCandidate]:
    document = Document(path)
    candidates: list[PlacementCandidate] = []
    for index, paragraph in enumerate(document.paragraphs):
        candidates.extend(_paragraph_candidates(paragraph, index, "document"))
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if cell.text.strip():
                    continue
                neighbors = []
                if cell_index > 0:
                    neighbors.append(row.cells[cell_index - 1].text.strip())
                if row_index > 0:
                    neighbors.append(table.rows[row_index - 1].cells[cell_index].text.strip())
                label = next((item for item in neighbors if item), f"表格字段{table_index + 1}-{row_index + 1}-{cell_index + 1}")
                context = " | ".join(item.text for item in row.cells)
                candidates.append(PlacementCandidate(
                    kind="table_cell_fill", label=label.rstrip("：:"), table_index=table_index,
                    row_index=row_index, cell_index=cell_index, context=context,
                    fingerprint=_fingerprint("document", f"t{table_index}:r{row_index}:c{cell_index}", context),
                ))
    for section_index, section in enumerate(document.sections):
        for part_name, part in ((f"header:{section_index}", section.header), (f"footer:{section_index}", section.footer)):
            for index, paragraph in enumerate(part.paragraphs):
                candidates.extend(_paragraph_candidates(paragraph, index, part_name))
    return candidates


def _field_key(label: str, index: int) -> str:
    if re.fullmatch(r"[A-Za-z][A-Za-z0-9_]*", label):
        return label
    known = {
        "姓名": "name", "员工姓名": "employee_name", "日期": "date",
        "处罚": "action", "处罚方式": "action", "单位名称": "unit_name",
    }
    for chinese, key in known.items():
        if chinese in label:
            return key
    return f"field_{index}"


def compile_docx_template(
    path: str, infer_semantics: Callable | None = None
) -> CompiledTemplate:
    candidates = extract_docx_candidates(path)
    inferred = infer_semantics(candidates) if infer_semantics else None
    fields: list[CompiledField] = []
    used: set[str] = set()
    for index, candidate in enumerate(candidates, start=1):
        semantics = inferred[index - 1] if inferred and index <= len(inferred) else {}
        key = semantics.get("key") or _field_key(candidate.label, index)
        while key in used:
            key = f"{key}_{index}"
        used.add(key)
        value_type = semantics.get("value_type") or (
            "single_choice" if candidate.kind == "checkbox_select" else
            "date" if candidate.kind == "date_parts" else "text"
        )
        fields.append(CompiledField(
            key=key,
            label=semantics.get("label") or candidate.label,
            value_type=value_type,
            options=candidate.options,
            required=bool(semantics.get("required", False)),
            fill_source=semantics.get("fill_source", "ai_then_user"),
            placements=[Placement(**candidate.model_dump(exclude={"label", "options"}), option_marks={option: index for index, option in enumerate(candidate.options)})],
        ))
    return CompiledTemplate(
        kind="docx", title=Path(path).stem, fields=fields,
        metadata={"candidate_count": len(candidates)},
        warnings=[] if candidates else ["未识别到可填写位置，请人工添加字段"],
    )
