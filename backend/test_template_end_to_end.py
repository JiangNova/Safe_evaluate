"""Regression tests covering representative business-document structures."""

import os
import tempfile
import unittest

from docx import Document

from backend.docx_renderer import render_compiled_docx
from backend.docx_template_compiler import compile_docx_template
from backend.document_quality import validate_document_for_finalize
from backend.template_ir import CompiledField


class UniversalTemplateEndToEndTests(unittest.TestCase):
    def test_employee_discipline_form_fills_table_and_checkbox(self):
        with tempfile.TemporaryDirectory() as directory:
            template = os.path.join(directory, "员工处罚单.docx")
            output = os.path.join(directory, "result.docx")
            document = Document()
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "员工姓名"
            table.cell(0, 1).text = ""
            document.add_paragraph("处罚方式：□警告  □记过  □辞退")
            document.save(template)
            compiled = compile_docx_template(template)
            table_field = next(field for field in compiled.fields if field.placements[0].kind == "table_cell_fill")
            choice_field = next(field for field in compiled.fields if field.placements[0].kind == "checkbox_select")
            render_compiled_docx(template, compiled, {
                table_field.key: {"value": "张三"},
                choice_field.key: {"value": "记过"},
            }, output, draft=False)
            rendered = Document(output)
        self.assertEqual(rendered.tables[0].cell(0, 1).text, "张三")
        self.assertIn("☑记过", rendered.paragraphs[0].text)

    def test_required_signature_blocks_formal_finalize(self):
        with tempfile.TemporaryDirectory() as directory:
            template = os.path.join(directory, "员工处分决定.docx")
            document = Document(); document.add_paragraph("签名：{{signature}}"); document.save(template)
            compiled = compile_docx_template(template)
        original = compiled.fields[0]
        compiled.fields[0] = original.model_copy(update={
            "fill_source": "user", "required": True, "missing_policy": "block_finalize"
        })
        report = validate_document_for_finalize(compiled, {"signature": {"value": ""}}, [])
        self.assertFalse(report.can_finalize)


if __name__ == "__main__":
    unittest.main()
