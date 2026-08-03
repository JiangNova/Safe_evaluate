"""HTTP contract tests for the anonymous leadership writing API."""

import io
import json
import os
import unittest
from unittest.mock import AsyncMock, patch

from docx import Document
from fastapi.testclient import TestClient

from backend.leadership_writer import GeneratedDocument, LeadershipWriterError
from backend.main import app


client = TestClient(app)


def leadership_token(username: str = "wanxin", password: str = "wanxin") -> str:
    response = client.post(
        "/api/leader-assistant/auth/login",
        json={"username": username, "password": password},
    )
    assert response.status_code == 200
    return response.json()["token"]


def leadership_headers() -> dict:
    return {"Authorization": f"Bearer {leadership_token()}"}


def profile_payload() -> dict:
    return {
        "name": "化学学院党委书记",
        "title": "党委书记",
        "organization": "化学学院",
        "responsibilities": "党建和安全工作",
        "focus_areas": "实验室危化品安全",
        "writing_preferences": "正式、简洁",
        "notes": "",
    }


def generated_document() -> GeneratedDocument:
    return GeneratedDocument(
        title="工作部署",
        content_markdown="# 工作部署\n\n请核实：具体完成时限。",
        warnings=["请核实具体完成时限。"],
    )


class LeadershipRouteTests(unittest.TestCase):
    def test_login_accepts_only_configured_leadership_accounts(self):
        for username, password in (("wanxin", "wanxin"), ("wanqin", "wanqin")):
            response = client.post(
                "/api/leader-assistant/auth/login",
                json={"username": username, "password": password},
            )
            self.assertEqual(response.status_code, 200)
            self.assertEqual(response.json()["user"], {"username": username, "role": "leader_assistant"})

        rejected = client.post(
            "/api/leader-assistant/auth/login",
            json={"username": "wanxin", "password": "incorrect"},
        )
        self.assertEqual(rejected.status_code, 401)

    def test_export_requires_leadership_token(self):
        response = client.post(
            "/api/leader-assistant/export/docx",
            json={"title": "工作部署", "content_markdown": "正文"},
        )
        self.assertEqual(response.status_code, 401)

    def test_generate_returns_document_without_persisting_upload(self):
        with patch(
            "backend.leadership_routes.generate_document",
            new=AsyncMock(return_value=generated_document()),
        ) as mocked_generate:
            response = client.post(
                "/api/leader-assistant/generate",
                headers=leadership_headers(),
                data={
                    "profile": json.dumps(profile_payload(), ensure_ascii=False),
                    "task_type": "safety_deployment",
                    "requirement": "根据材料起草安全工作部署。",
                },
                files={"files": ("notes.txt", "落实实验室安全检查", "text/plain")},
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["title"], "工作部署")
        sources = mocked_generate.await_args.args[2]
        self.assertEqual(sources[0].filename, "notes.txt")

    def test_generate_rejects_invalid_profile_before_writer_call(self):
        response = client.post(
            "/api/leader-assistant/generate",
            headers=leadership_headers(),
            data={
                "profile": "{}",
                "task_type": "safety_deployment",
                "requirement": "起草部署",
            },
        )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["detail"]["code"], "invalid_request")
        self.assertEqual(response.json()["detail"]["stage"], "validation")

    def test_generate_returns_provider_error(self):
        with patch(
            "backend.leadership_routes.generate_document",
            new=AsyncMock(side_effect=RuntimeError("provider unavailable")),
        ):
            response = client.post(
                "/api/leader-assistant/generate",
                headers=leadership_headers(),
                data={
                    "profile": json.dumps(profile_payload(), ensure_ascii=False),
                    "task_type": "summary",
                    "requirement": "起草总结",
                },
            )

        self.assertEqual(response.status_code, 502)
        self.assertEqual(response.json()["detail"]["stage"], "generation")

    def test_generate_cleans_temporary_upload_after_extraction(self):
        seen_paths: list[str] = []

        def observe_source(record):
            seen_paths.append(record["storage_path"])
            self.assertTrue(os.path.exists(record["storage_path"]))
            from backend.public_files import ParsedSource, SourceChunk

            return ParsedSource(
                file_id=record["id"],
                filename=record["original_name"],
                media_type="text",
                chunks=[SourceChunk("参考材料", "text:1")],
                warnings=[],
            )

        with (
            patch("backend.leadership_routes.extract_source", side_effect=observe_source),
            patch(
                "backend.leadership_routes.generate_document",
                new=AsyncMock(return_value=generated_document()),
            ),
        ):
            response = client.post(
                "/api/leader-assistant/generate",
                headers=leadership_headers(),
                data={
                    "profile": json.dumps(profile_payload(), ensure_ascii=False),
                    "task_type": "custom",
                    "requirement": "起草文稿",
                },
                files={"files": ("notes.txt", "参考材料", "text/plain")},
            )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(seen_paths)
        self.assertTrue(all(not os.path.exists(path) for path in seen_paths))

    def test_revise_returns_document(self):
        with patch(
            "backend.leadership_routes.revise_document",
            new=AsyncMock(return_value=generated_document()),
        ) as mocked_revise:
            response = client.post(
                "/api/leader-assistant/revise",
                headers=leadership_headers(),
                json={
                    "profile": profile_payload(),
                    "task_type": "implementation_report",
                    "requirement": "起草落实报告",
                    "title": "初稿",
                    "content_markdown": "# 初稿\n\n正文",
                    "warnings": [],
                    "revision_instruction": "增强安全职责表述",
                },
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(mocked_revise.await_args.args[3], "增强安全职责表述")

    def test_export_returns_docx_attachment(self):
        response = client.post(
            "/api/leader-assistant/export/docx",
            headers=leadership_headers(),
            json={"title": "工作部署", "content_markdown": "# 工作部署\n\n正文"},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("attachment", response.headers["content-disposition"])
        document = Document(io.BytesIO(response.content))
        self.assertIn("工作部署", "\n".join(p.text for p in document.paragraphs))
        self.assertIn("正文", "\n".join(p.text for p in document.paragraphs))


if __name__ == "__main__":
    unittest.main()
