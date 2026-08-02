"""Placement executor for compiled structural DOCX templates."""

from __future__ import annotations

import hashlib
import re
from copy import deepcopy

from docx import Document

from .document_renderer import RenderResult, RenderWarning
from .template_ir import CompiledTemplate, Placement


MISSING_MARKER = "[待人工补充]"


def _text_value(values: dict, key: str) -> str:
    raw = values.get(key)
    if isinstance(raw, dict):
        raw = raw.get("value")
    if raw is None or raw == "":
        return ""
    if isinstance(raw, list):
        return "、".join(str(item) for item in raw)
    return str(raw)


def _paragraph_for(document, placement: Placement):
    if placement.part == "document":
        return document.paragraphs[placement.paragraph_index or 0]
    kind, section_index = placement.part.split(":", 1)
    section = document.sections[int(section_index)]
    paragraphs = section.header.paragraphs if kind == "header" else section.footer.paragraphs
    return paragraphs[placement.paragraph_index or 0]


def _current_fingerprint(placement: Placement, context: str) -> str:
    if placement.kind == "run_range_replace":
        path = f"p{placement.paragraph_index}:r{placement.run_start}"
    elif placement.kind == "date_parts":
        path = f"p{placement.paragraph_index}:date"
    elif placement.kind == "checkbox_select":
        path = f"p{placement.paragraph_index}:checkbox"
    elif placement.kind == "placeholder_replace":
        path = f"p{placement.paragraph_index}:placeholder:{(placement.placeholder or '').strip('{} ')}"
    elif placement.kind == "table_cell_fill":
        path = f"t{placement.table_index}:r{placement.row_index}:c{placement.cell_index}"
    else:
        return placement.fingerprint
    normalized = re.sub(r"\s+", " ", context).strip()
    return hashlib.sha256(f"{placement.part}|{path}|{normalized}".encode("utf-8")).hexdigest()


def _replace_all_runs(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _apply_paragraph_placement(paragraph, placement: Placement, value: str) -> None:
    if placement.kind == "run_range_replace":
        start = placement.run_start or 0
        end = placement.run_end if placement.run_end is not None else start
        paragraph.runs[start].text = value
        for index in range(start + 1, min(end + 1, len(paragraph.runs))):
            paragraph.runs[index].text = ""
    elif placement.kind == "placeholder_replace":
        _replace_all_runs(paragraph, paragraph.text.replace(placement.placeholder or "", value))
    elif placement.kind in {"paragraph_insert", "header_footer_fill"}:
        anchor = placement.anchor or ""
        if anchor and anchor in paragraph.text:
            _replace_all_runs(paragraph, paragraph.text.replace(anchor, f"{anchor}{value}", 1))
        else:
            paragraph.add_run(value)
    elif placement.kind == "date_parts":
        match = re.match(r"(\d{4})[-/.年](\d{1,2})[-/.月](\d{1,2})", value)
        parts = match.groups() if match else (MISSING_MARKER, MISSING_MARKER, MISSING_MARKER)
        rendered = re.sub(
            r"年\s*[_＿\s]{2,}\s*月\s*[_＿\s]{2,}\s*日",
            f"年 {parts[0]} 月 {parts[1]} 日 {parts[2]}", paragraph.text,
        )
        _replace_all_runs(paragraph, rendered)
    elif placement.kind == "checkbox_select":
        options = placement.option_marks or {}
        selected = {item.strip() for item in value.split("、") if item.strip()}
        rendered = paragraph.text
        for option in options:
            mark = "☑" if option in selected else "□"
            rendered = re.sub(rf"[□☐☑■](\s*){re.escape(option)}", rf"{mark}\1{option}", rendered)
        _replace_all_runs(paragraph, rendered)


def render_compiled_docx(
    template_path: str,
    compiled: CompiledTemplate,
    values: dict,
    output_path: str,
    *,
    draft: bool = True,
) -> RenderResult:
    document = Document(template_path)
    warnings: list[RenderWarning] = []
    for field in compiled.fields:
        value = _text_value(values, field.key)
        if not value and field.required:
            warnings.append(RenderWarning("missing_required_field", f"必填字段“{field.label}”为空", field.key))
        rendered_value = value or (MISSING_MARKER if draft and field.missing_policy == "block_finalize" else "")
        for placement in field.placements:
            try:
                if placement.kind == "table_cell_fill":
                    table = document.tables[placement.table_index or 0]
                    row = table.rows[placement.row_index or 0]
                    context = " | ".join(cell.text for cell in row.cells)
                    if placement.fingerprint and _current_fingerprint(placement, context) != placement.fingerprint:
                        raise ValueError("stale")
                    row.cells[placement.cell_index or 0].text = rendered_value
                elif placement.kind == "repeat_table_row":
                    table = document.tables[placement.table_index or 0]
                    source_row = table.rows[placement.row_index or 0]._tr
                    rows = values.get(field.key, {}).get("value", []) if isinstance(values.get(field.key), dict) else values.get(field.key, [])
                    for item in rows[1:]:
                        copied = deepcopy(source_row)
                        table._tbl.append(copied)
                        for cell, cell_value in zip(table.rows[-1].cells, item.values() if isinstance(item, dict) else item):
                            cell.text = str(cell_value)
                else:
                    paragraph = _paragraph_for(document, placement)
                    if placement.fingerprint and _current_fingerprint(placement, paragraph.text) != placement.fingerprint:
                        raise ValueError("stale")
                    _apply_paragraph_placement(paragraph, placement, rendered_value)
            except (IndexError, KeyError, ValueError):
                warnings.append(RenderWarning("stale_placement", f"字段“{field.label}”的填写位置已变化", field.key))
    document.save(output_path)
    return RenderResult(output_path, warnings)
