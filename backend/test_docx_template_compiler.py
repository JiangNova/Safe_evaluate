import os
import tempfile
import unittest

from docx import Document

from backend.docx_template_compiler import compile_docx_template, extract_docx_candidates


class DocxTemplateCompilerTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.path = os.path.join(self.temp.name, "员工处罚.docx")
        document = Document()
        paragraph = document.add_paragraph("员工：")
        blank = paragraph.add_run("________")
        blank.font.underline = True
        document.add_paragraph("日期：年    月    日")
        document.add_paragraph("处罚：□警告  □记过  □辞退")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "事实描述"
        table.cell(0, 1).text = ""
        document.save(self.path)

    def tearDown(self):
        self.temp.cleanup()

    def test_extracts_blanks_date_parts_checkboxes_and_table_cells(self):
        kinds = {item.kind for item in extract_docx_candidates(self.path)}
        self.assertTrue({"run_range_replace", "date_parts", "checkbox_select", "table_cell_fill"}.issubset(kinds))

    def test_compiles_candidates_to_typed_fields(self):
        compiled = compile_docx_template(self.path)
        self.assertEqual(compiled.kind, "docx")
        self.assertGreaterEqual(len(compiled.fields), 4)
        self.assertEqual(next(field for field in compiled.fields if field.value_type == "single_choice").options, ["警告", "记过", "辞退"])


if __name__ == "__main__":
    unittest.main()
