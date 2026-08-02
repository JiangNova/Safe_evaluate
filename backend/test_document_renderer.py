"""Tests for deterministic template rendering and artifact packaging."""

import importlib.util
import io
import os
import tempfile
import unittest
import zipfile

from docx import Document

from backend import document_renderer
from backend.template_parser import TemplateField


HAS_PDF_DEPS = bool(
    importlib.util.find_spec("pypdf") and importlib.util.find_spec("reportlab")
)


def all_docx_text(path: str) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.header.paragraphs)
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(parts)


class DocumentRendererTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_render_docx_replaces_split_body_table_header_and_footer_fields(self):
        template_path = os.path.join(self.temp_dir.name, "template.docx")
        output_path = os.path.join(self.temp_dir.name, "output.docx")
        document = Document()
        paragraph = document.add_paragraph("单位：")
        paragraph.add_run("{{unit_")
        paragraph.add_run("name}}")
        document.add_table(rows=1, cols=1).cell(0, 0).text = "{{summary}}"
        document.sections[0].header.paragraphs[0].text = "{{document_number}}"
        document.sections[0].footer.paragraphs[0].text = "{{unit_name}}"
        document.save(template_path)
        fields = [
            self._docx_field("unit_name"),
            self._docx_field("summary", "multiline"),
            self._docx_field("document_number"),
        ]
        values = {
            "unit_name": {"value": "示例公司"},
            "summary": {"value": "第一项\n第二项"},
            "document_number": {"value": "2026-001"},
        }

        result = document_renderer.render_docx(
            template_path, fields, values, output_path
        )
        rendered = all_docx_text(result.path)

        self.assertNotIn("{{", rendered)
        self.assertGreaterEqual(rendered.count("示例公司"), 2)
        self.assertIn("第一项\n第二项", rendered)
        self.assertIn("2026-001", rendered)

    def test_missing_required_value_returns_warning(self):
        template_path = os.path.join(self.temp_dir.name, "template.docx")
        output_path = os.path.join(self.temp_dir.name, "output.docx")
        document = Document()
        document.add_paragraph("{{required_field}}")
        document.save(template_path)

        result = document_renderer.render_docx(
            template_path,
            [self._docx_field("required_field", required=True)],
            {},
            output_path,
        )

        self.assertIn("missing_required_field", [item.code for item in result.warnings])

    def test_render_docx_fills_inferred_field_after_anchor(self):
        template_path = os.path.join(self.temp_dir.name, "inferred-template.docx")
        output_path = os.path.join(self.temp_dir.name, "inferred-output.docx")
        document = Document()
        document.add_paragraph("Unit name: ")
        document.save(template_path)
        field = TemplateField(
            key="unit_name",
            label="Unit name",
            field_type="text",
            required=True,
            repeating=False,
            confidence=0.82,
            locator={"kind": "docx_inferred", "anchor": "Unit name: "},
        )

        result = document_renderer.render_docx(
            template_path,
            [field],
            {"unit_name": {"value": "Example Company"}},
            output_path,
        )

        self.assertIn("Unit name: Example Company", all_docx_text(result.path))
        self.assertNotIn("anchor_not_found", [item.code for item in result.warnings])

    @unittest.skipUnless(HAS_PDF_DEPS, "PDF runtime dependencies are not installed locally")
    def test_render_pdf_preserves_page_count_and_reports_overflow(self):
        from reportlab.pdfgen import canvas
        from pypdf import PdfReader

        template_path = os.path.join(self.temp_dir.name, "template.pdf")
        output_path = os.path.join(self.temp_dir.name, "output.pdf")
        pdf = canvas.Canvas(template_path, pagesize=(300, 300))
        pdf.drawString(20, 270, "Template")
        pdf.save()
        field = TemplateField(
            key="summary",
            label="Summary",
            field_type="multiline",
            required=True,
            repeating=False,
            confidence=1.0,
            locator={"kind": "pdf_rect", "page": 0, "rect": [20, 20, 100, 40]},
        )

        result = document_renderer.render_pdf(
            template_path,
            [field],
            {"summary": {"value": "very long text " * 50}},
            output_path,
        )

        self.assertEqual(len(PdfReader(result.path).pages), 1)
        self.assertIn("field_overflow", [item.code for item in result.warnings])

    def test_zip_contains_outputs_and_failure_manifest(self):
        output_file = os.path.join(self.temp_dir.name, "report.docx")
        with open(output_file, "wb") as stream:
            stream.write(b"docx")
        archive_path = os.path.join(self.temp_dir.name, "outputs.zip")

        result = document_renderer.build_artifact_zip(
            [(output_file, "report.docx")],
            [{"template": "form.pdf", "error": "render failed"}],
            archive_path,
        )

        with zipfile.ZipFile(result) as archive:
            self.assertEqual(
                set(archive.namelist()), {"report.docx", "失败清单.txt"}
            )
            self.assertIn(
                "form.pdf", archive.read("失败清单.txt").decode("utf-8")
            )

    @staticmethod
    def _docx_field(key, field_type="text", required=False):
        return TemplateField(
            key=key,
            label=key,
            field_type=field_type,
            required=required,
            repeating=False,
            confidence=1.0,
            locator={"kind": "docx_placeholder", "locations": []},
        )


if __name__ == "__main__":
    unittest.main()
