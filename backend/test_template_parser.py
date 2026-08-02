"""Tests for DOCX/PDF output template parsing."""

import os
import tempfile
import unittest
from unittest.mock import patch

from docx import Document

from backend import template_parser


class TemplateParserTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()

    def tearDown(self):
        self.temp_dir.cleanup()

    def _save_docx_template(self) -> dict:
        path = os.path.join(self.temp_dir.name, "template.docx")
        document = Document()
        paragraph = document.add_paragraph("单位名称：")
        paragraph.add_run("{{unit_")
        paragraph.add_run("name}}")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "检查日期：{{inspection_date}}"
        header = document.sections[0].header.paragraphs[0]
        header.text = "编号：{{document_number}}"
        footer = document.sections[0].footer.paragraphs[0]
        footer.text = "{{unit_name}}"
        document.save(path)
        return {
            "id": 1,
            "safe_name": "template.docx",
            "original_name": "检查表.docx",
            "storage_path": path,
        }

    def test_docx_parser_finds_split_body_table_header_and_footer_placeholders(self):
        result = template_parser.parse_template(self._save_docx_template())

        self.assertEqual(
            set(result.fields_by_key),
            {"unit_name", "inspection_date", "document_number"},
        )
        unit_locations = result.fields_by_key["unit_name"].locator["locations"]
        self.assertEqual(len(unit_locations), 2)
        self.assertEqual(
            {location["container"] for location in unit_locations},
            {"body", "footer"},
        )
        self.assertFalse(result.requires_confirmation)

    def test_docx_without_placeholders_uses_inference_callback(self):
        path = os.path.join(self.temp_dir.name, "plain.docx")
        document = Document()
        document.add_paragraph("单位名称：")
        document.save(path)
        record = {
            "id": 2,
            "safe_name": "plain.docx",
            "original_name": "plain.docx",
            "storage_path": path,
        }

        result = template_parser.parse_template(
            record,
            infer_fields=lambda text, layout: [
                {
                    "key": "unit_name",
                    "label": "单位名称",
                    "field_type": "text",
                    "required": True,
                    "repeating": False,
                    "confidence": 0.83,
                    "locator": {"kind": "docx_inferred", "anchor": "单位名称："},
                }
            ],
        )

        self.assertEqual(result.fields[0].key, "unit_name")
        self.assertTrue(result.requires_confirmation)

    def test_pdf_parser_emits_page_rect_and_requires_confirmation(self):
        path = os.path.join(self.temp_dir.name, "template.pdf")
        with open(path, "wb") as output:
            output.write(b"%PDF-1.7\n%%EOF")
        record = {
            "id": 3,
            "safe_name": "template.pdf",
            "original_name": "template.pdf",
            "storage_path": path,
        }
        layout = [
            {
                "page": 0,
                "width": 595.0,
                "height": 842.0,
                "blocks": [
                    {"text": "单位名称 {{unit_name}}", "rect": [100, 700, 350, 730]}
                ],
            }
        ]

        with patch.object(template_parser, "_read_pdf_layout", return_value=layout):
            result = template_parser.parse_template(record)

        field = result.fields_by_key["unit_name"]
        self.assertEqual(field.locator["page"], 0)
        self.assertEqual(len(field.locator["rect"]), 4)
        self.assertTrue(result.requires_confirmation)

    def test_overlapping_pdf_rectangles_are_rejected(self):
        fields = [
            {
                "key": "field_a",
                "label": "A",
                "field_type": "text",
                "required": False,
                "repeating": False,
                "confidence": 1.0,
                "locator": {"kind": "pdf_rect", "page": 0, "rect": [10, 10, 100, 40]},
            },
            {
                "key": "field_b",
                "label": "B",
                "field_type": "text",
                "required": False,
                "repeating": False,
                "confidence": 1.0,
                "locator": {"kind": "pdf_rect", "page": 0, "rect": [50, 20, 120, 50]},
            },
        ]

        with self.assertRaises(template_parser.TemplateFieldError):
            template_parser.validate_field_definitions("pdf", fields)


if __name__ == "__main__":
    unittest.main()
