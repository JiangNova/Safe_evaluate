"""Secure storage and source extraction for anonymous public job files."""

from __future__ import annotations

import io
import os
import re
import secrets
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from .config import (
    MAX_FILE_SIZE,
    PUBLIC_JOB_MAX_FILES,
    PUBLIC_JOB_MAX_TOTAL_SIZE,
    PUBLIC_JOB_STORAGE_DIR,
)
from . import public_jobs


DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

ALLOWED_BY_KIND = {
    "material": {".png", ".jpg", ".jpeg", ".webp", ".pdf", ".docx"},
    "basis": {".pdf", ".docx", ".txt"},
    "template": {".pdf", ".docx"},
}

EXPECTED_MIMES = {
    ".png": {"image/png", "application/octet-stream"},
    ".jpg": {"image/jpeg", "application/octet-stream"},
    ".jpeg": {"image/jpeg", "application/octet-stream"},
    ".webp": {"image/webp", "application/octet-stream"},
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".docx": {DOCX_MIME, "application/zip", "application/octet-stream"},
    ".txt": {"text/plain", "application/octet-stream"},
}

MAX_DOCX_ENTRIES = 1000
MAX_DOCX_UNCOMPRESSED_SIZE = 100 * 1024 * 1024


class UploadValidationError(ValueError):
    """The uploaded file violates format or safety constraints."""


class SourceExtractionError(RuntimeError):
    """A validated file could not be converted into evaluation sources."""


@dataclass(frozen=True)
class ValidatedUpload:
    kind: str
    original_name: str
    mime_type: str
    extension: str
    data: bytes


@dataclass(frozen=True)
class SourceChunk:
    text: str
    source_ref: str


@dataclass(frozen=True)
class ParsedSource:
    file_id: int
    filename: str
    media_type: str
    chunks: list[SourceChunk]
    warnings: list[str]


def _clean_original_name(filename: str) -> str:
    name = os.path.basename((filename or "").replace("\\", "/")).strip()
    if not name or name in {".", ".."}:
        raise UploadValidationError("文件名无效")
    return name[:255]


def _validate_docx(data: bytes) -> None:
    if not data.startswith(b"PK"):
        raise UploadValidationError("DOCX 文件签名无效")
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
            names = {info.filename for info in infos}
            if len(infos) > MAX_DOCX_ENTRIES:
                raise UploadValidationError("DOCX 内部文件数量超限")
            total = sum(info.file_size for info in infos)
            if total > MAX_DOCX_UNCOMPRESSED_SIZE:
                raise UploadValidationError("DOCX 解压后大小超限")
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise UploadValidationError("文件不是有效的 DOCX 文档")
            if any(
                name.lower().endswith(("vbaproject.bin", ".exe", ".dll"))
                for name in names
            ):
                raise UploadValidationError("DOCX 包含不允许的可执行或宏内容")
    except zipfile.BadZipFile as exc:
        raise UploadValidationError("DOCX 文件已损坏") from exc


def _validate_signature(extension: str, data: bytes) -> None:
    if extension == ".pdf":
        if not data.startswith(b"%PDF-"):
            raise UploadValidationError("PDF 文件签名无效")
        if re.search(br"/Encrypt\b", data):
            raise UploadValidationError("暂不支持加密 PDF")
    elif extension == ".docx":
        _validate_docx(data)
    elif extension == ".png" and not data.startswith(b"\x89PNG\r\n\x1a\n"):
        raise UploadValidationError("PNG 文件签名无效")
    elif extension in {".jpg", ".jpeg"} and not data.startswith(b"\xff\xd8\xff"):
        raise UploadValidationError("JPEG 文件签名无效")
    elif extension == ".webp" and not (
        data.startswith(b"RIFF") and data[8:12] == b"WEBP"
    ):
        raise UploadValidationError("WebP 文件签名无效")


def validate_upload(
    kind: str, filename: str, mime: str, data: bytes
) -> ValidatedUpload:
    """Validate an upload using its declared kind, suffix, MIME, and magic."""
    if kind not in ALLOWED_BY_KIND:
        raise UploadValidationError("不支持的文件用途")
    original_name = _clean_original_name(filename)
    extension = Path(original_name).suffix.lower()
    if extension not in ALLOWED_BY_KIND[kind]:
        raise UploadValidationError(f"{kind} 不支持 {extension or '无扩展名'} 文件")
    if not data:
        raise UploadValidationError("文件为空")
    if len(data) > MAX_FILE_SIZE:
        raise UploadValidationError("单个文件超过 50MB 限制")
    normalized_mime = (mime or "application/octet-stream").lower()
    if normalized_mime not in EXPECTED_MIMES[extension]:
        raise UploadValidationError("文件 MIME 与扩展名不匹配")
    _validate_signature(extension, data)
    return ValidatedUpload(
        kind=kind,
        original_name=original_name,
        mime_type=normalized_mime,
        extension=extension,
        data=data,
    )


def store_upload(job_id: str, upload: ValidatedUpload) -> dict:
    """Store validated bytes in an isolated job directory and persist metadata."""
    file_count, total_size = public_jobs.get_file_usage(job_id)
    if file_count >= PUBLIC_JOB_MAX_FILES:
        raise UploadValidationError("本次任务文件数量已达上限")
    if total_size + len(upload.data) > PUBLIC_JOB_MAX_TOTAL_SIZE:
        raise UploadValidationError("本次任务文件总大小已达上限")

    job_dir = os.path.abspath(os.path.join(PUBLIC_JOB_STORAGE_DIR, job_id))
    storage_root = os.path.abspath(PUBLIC_JOB_STORAGE_DIR)
    if os.path.commonpath([storage_root, job_dir]) != storage_root:
        raise UploadValidationError("任务存储路径无效")
    os.makedirs(job_dir, exist_ok=True)

    safe_name = f"{upload.kind}-{secrets.token_hex(12)}{upload.extension}"
    storage_path = os.path.join(job_dir, safe_name)
    with open(storage_path, "xb") as output:
        output.write(upload.data)

    try:
        return public_jobs.add_file(
            job_id,
            upload.kind,
            {
                "safe_name": safe_name,
                "original_name": upload.original_name,
                "mime_type": upload.mime_type,
                "size": len(upload.data),
                "storage_path": storage_path,
                "parse_status": "pending",
                "parse_metadata_json": None,
            },
        )
    except Exception:
        try:
            os.remove(storage_path)
        except OSError:
            pass
        raise


def _read_text(path: str) -> tuple[str, list[str]]:
    data = Path(path).read_bytes()
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding), []
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), ["文本编码无法完全识别，已替换异常字符"]


def _read_docx_chunks(path: str) -> list[SourceChunk]:
    document = Document(path)
    chunks: list[SourceChunk] = []
    paragraph_index = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraph_index += 1
            chunks.append(SourceChunk(text, f"paragraph:{paragraph_index}"))
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            text = " | ".join(cell.text.strip() for cell in row.cells).strip(" |")
            if text:
                chunks.append(
                    SourceChunk(text, f"table:{table_index}:row:{row_index}")
                )
    return chunks


def _read_pdf_pages(path: str) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise SourceExtractionError("缺少 pypdf，无法解析 PDF") from exc
    try:
        reader = PdfReader(path)
        if reader.is_encrypted:
            raise SourceExtractionError("暂不支持加密 PDF")
        return [(page.extract_text() or "").strip() for page in reader.pages]
    except SourceExtractionError:
        raise
    except Exception as exc:
        raise SourceExtractionError(f"PDF 解析失败: {exc}") from exc


def extract_source(file_record: dict) -> ParsedSource:
    """Extract stable source chunks while retaining file/page references."""
    extension = Path(file_record["safe_name"]).suffix.lower()
    chunks: list[SourceChunk]
    warnings: list[str] = []
    media_type = "text"

    if extension == ".docx":
        chunks = _read_docx_chunks(file_record["storage_path"])
    elif extension == ".pdf":
        pages = _read_pdf_pages(file_record["storage_path"])
        chunks = [
            SourceChunk(text, f"page:{index}")
            for index, text in enumerate(pages, start=1)
            if text
        ]
        if not chunks:
            warnings.append("PDF 未提取到文本，后续需要 OCR 或视觉模型识别")
    elif extension == ".txt":
        text, warnings = _read_text(file_record["storage_path"])
        chunks = [SourceChunk(text.strip(), "text:1")] if text.strip() else []
    elif extension in {".png", ".jpg", ".jpeg", ".webp"}:
        media_type = "image"
        chunks = [
            SourceChunk(
                f"[视觉材料: {file_record['original_name']}]",
                f"image:{file_record['original_name']}",
            )
        ]
    else:
        raise SourceExtractionError(f"无法解析的文件格式: {extension}")

    return ParsedSource(
        file_id=int(file_record["id"]),
        filename=file_record["original_name"],
        media_type=media_type,
        chunks=chunks,
        warnings=warnings,
    )
