"""Parse DOCX/PDF output templates into user-confirmable field definitions."""

from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Callable, Iterable, Literal

from docx import Document


PLACEHOLDER_RE = re.compile(
    r"\{\{\s*([A-Za-z][A-Za-z0-9_.-]{0,63})\s*\}\}"
)
FIELD_KEY_RE = re.compile(r"^[A-Za-z][A-Za-z0-9_.-]{0,63}$")
FIELD_TYPES = {"text", "multiline", "date", "boolean", "list"}


class TemplateFieldError(ValueError):
    """A field definition is invalid or ambiguous."""


@dataclass(frozen=True)
class TemplateField:
    key: str
    label: str
    field_type: Literal["text", "multiline", "date", "boolean", "list"]
    required: bool
    repeating: bool
    confidence: float
    locator: dict

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass(frozen=True)
class TemplateParseResult:
    source_format: str
    fields: list[TemplateField]
    preview_metadata: dict
    warnings: list[str]
    requires_confirmation: bool

    @property
    def fields_by_key(self) -> dict[str, TemplateField]:
        return {field.key: field for field in self.fields}

    def to_dict(self) -> dict:
        return {
            "source_format": self.source_format,
            "fields": [field.to_dict() for field in self.fields],
            "preview_metadata": self.preview_metadata,
            "warnings": self.warnings,
            "requires_confirmation": self.requires_confirmation,
        }


InferenceCallback = Callable[[str, list[dict]], list[dict]]


def _iter_table_paragraphs(table, table_path: str):
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell_path = f"{table_path}.row:{row_index}.cell:{cell_index}"
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                yield "table", f"{cell_path}.paragraph:{paragraph_index}", paragraph
            for nested_index, nested in enumerate(cell.tables):
                yield from _iter_table_paragraphs(
                    nested, f"{cell_path}.table:{nested_index}"
                )


def _iter_docx_paragraphs(document: Document):
    for index, paragraph in enumerate(document.paragraphs):
        yield "body", f"paragraph:{index}", paragraph
    for table_index, table in enumerate(document.tables):
        yield from _iter_table_paragraphs(table, f"table:{table_index}")
    for section_index, section in enumerate(document.sections):
        for container_name, part in (
            ("header", section.header),
            ("footer", section.footer),
        ):
            for index, paragraph in enumerate(part.paragraphs):
                yield (
                    container_name,
                    f"section:{section_index}.{container_name}.paragraph:{index}",
                    paragraph,
                )
            for table_index, table in enumerate(part.tables):
                for _, path, paragraph in _iter_table_paragraphs(
                    table,
                    f"section:{section_index}.{container_name}.table:{table_index}",
                ):
                    yield container_name, path, paragraph


def _label_for_key(key: str) -> str:
    return key.replace("_", " ").replace(".", " ").strip().title()


def _parse_docx(path: str, infer_fields: InferenceCallback | None) -> TemplateParseResult:
    document = Document(path)
    locations_by_key: dict[str, list[dict]] = {}
    text_parts: list[str] = []
    layout: list[dict] = []

    for container, location_path, paragraph in _iter_docx_paragraphs(document):
        text = paragraph.text
        if text.strip():
            text_parts.append(text.strip())
            layout.append(
                {
                    "container": container,
                    "path": location_path,
                    "text": text,
                }
            )
        for match in PLACEHOLDER_RE.finditer(text):
            locations_by_key.setdefault(match.group(1), []).append(
                {
                    "container": container,
                    "path": location_path,
                    "placeholder": match.group(0),
                }
            )

    if locations_by_key:
        fields = [
            TemplateField(
                key=key,
                label=_label_for_key(key),
                field_type="text",
                required=False,
                repeating=False,
                confidence=1.0,
                locator={"kind": "docx_placeholder", "locations": locations},
            )
            for key, locations in locations_by_key.items()
        ]
        return TemplateParseResult(
            source_format="docx",
            fields=fields,
            preview_metadata={"paragraphs": layout},
            warnings=[],
            requires_confirmation=False,
        )

    inferred = infer_fields("\n".join(text_parts), layout) if infer_fields else []
    fields = validate_field_definitions("docx", inferred) if inferred else []
    warnings = [] if fields else ["模板未包含占位符，也未自动识别到可填写字段"]
    return TemplateParseResult(
        source_format="docx",
        fields=fields,
        preview_metadata={"paragraphs": layout},
        warnings=warnings,
        requires_confirmation=True,
    )


def _read_pdf_layout(path: str) -> list[dict]:
    """Extract PDF text fragments with approximate PDF-coordinate rectangles."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise TemplateFieldError("缺少 pypdf，无法解析 PDF 模板") from exc

    reader = PdfReader(path)
    if reader.is_encrypted:
        raise TemplateFieldError("暂不支持加密 PDF 模板")
    pages: list[dict] = []
    for page_index, page in enumerate(reader.pages):
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        blocks: list[dict] = []

        def visitor(text, cm, tm, font_dict, font_size):
            cleaned = (text or "").strip()
            if not cleaned:
                return
            x = float(tm[4])
            baseline = float(tm[5])
            estimated_width = max(float(font_size), len(cleaned) * float(font_size) * 0.55)
            blocks.append(
                {
                    "text": cleaned,
                    "rect": [
                        x,
                        max(0.0, baseline - float(font_size) * 0.25),
                        min(width, x + estimated_width),
                        min(height, baseline + float(font_size)),
                    ],
                }
            )

        page.extract_text(visitor_text=visitor)
        pages.append(
            {
                "page": page_index,
                "width": width,
                "height": height,
                "blocks": blocks,
            }
        )
    return pages


def _ocr_pdf_layout(path: str) -> list[dict]:
    """Render and OCR a scanned PDF when optional production dependencies exist."""
    try:
        import pypdfium2 as pdfium
        import pytesseract
        from pytesseract import Output
    except ImportError:
        return []

    document = pdfium.PdfDocument(path)
    pages: list[dict] = []
    scale = 200 / 72
    for page_index in range(len(document)):
        page = document[page_index]
        width, height = page.get_size()
        image = page.render(scale=scale).to_pil()
        data = pytesseract.image_to_data(
            image, lang="chi_sim+eng", output_type=Output.DICT
        )
        blocks: list[dict] = []
        for index, text in enumerate(data["text"]):
            cleaned = text.strip()
            if not cleaned:
                continue
            left = data["left"][index] / scale
            top = data["top"][index] / scale
            box_width = data["width"][index] / scale
            box_height = data["height"][index] / scale
            blocks.append(
                {
                    "text": cleaned,
                    "rect": [
                        left,
                        height - top - box_height,
                        left + box_width,
                        height - top,
                    ],
                }
            )
        pages.append(
            {
                "page": page_index,
                "width": width,
                "height": height,
                "blocks": blocks,
                "ocr": True,
            }
        )
    return pages


def _pdf_placeholder_field(page: dict, block: dict, match: re.Match) -> dict:
    x0, y0, x1, y1 = [float(value) for value in block["rect"]]
    text = block["text"]
    denominator = max(1, len(text))
    start_ratio = match.start() / denominator
    end_ratio = match.end() / denominator
    field_x0 = x0 + (x1 - x0) * start_ratio
    field_x1 = x0 + (x1 - x0) * end_ratio
    return {
        "key": match.group(1),
        "label": _label_for_key(match.group(1)),
        "field_type": "text",
        "required": False,
        "repeating": False,
        "confidence": 0.9,
        "locator": {
            "kind": "pdf_rect",
            "page": int(page["page"]),
            "rect": [field_x0, y0, max(field_x1, field_x0 + 12), y1],
        },
    }


def _parse_pdf(path: str, infer_fields: InferenceCallback | None) -> TemplateParseResult:
    layout = _read_pdf_layout(path)
    warnings: list[str] = []
    if not any(page.get("blocks") for page in layout):
        ocr_layout = _ocr_pdf_layout(path)
        if ocr_layout:
            layout = ocr_layout
            warnings.append("扫描型 PDF 已使用 OCR 识别，所有字段都需要人工确认")
        else:
            warnings.append("PDF 未提取到文本，当前环境无法执行 OCR")

    fields: list[dict] = []
    text_parts: list[str] = []
    for page in layout:
        for block in page.get("blocks", []):
            text = str(block.get("text", ""))
            if text:
                text_parts.append(f"[page:{int(page['page']) + 1}] {text}")
            fields.extend(
                _pdf_placeholder_field(page, block, match)
                for match in PLACEHOLDER_RE.finditer(text)
            )

    if not fields and infer_fields:
        fields = infer_fields("\n".join(text_parts), layout)
    validated = validate_field_definitions("pdf", fields) if fields else []
    if not validated:
        warnings.append("模板未自动识别到可填写字段")
    return TemplateParseResult(
        source_format="pdf",
        fields=validated,
        preview_metadata={"pages": layout},
        warnings=warnings,
        requires_confirmation=True,
    )


def _rectangles_overlap(first: list[float], second: list[float]) -> bool:
    return (
        min(first[2], second[2]) > max(first[0], second[0])
        and min(first[3], second[3]) > max(first[1], second[1])
    )


def validate_field_definitions(
    source_format: str, fields: Iterable[dict | TemplateField]
) -> list[TemplateField]:
    """Normalize user/AI field definitions and reject unsafe ambiguity."""
    normalized: list[TemplateField] = []
    seen_keys: set[str] = set()
    for raw in fields:
        data = raw.to_dict() if isinstance(raw, TemplateField) else dict(raw)
        key = str(data.get("key", "")).strip()
        if not FIELD_KEY_RE.fullmatch(key):
            raise TemplateFieldError(f"字段键无效: {key}")
        if key in seen_keys:
            raise TemplateFieldError(f"字段键重复: {key}")
        seen_keys.add(key)
        field_type = str(data.get("field_type", "text"))
        if field_type not in FIELD_TYPES:
            raise TemplateFieldError(f"字段类型无效: {field_type}")
        confidence = float(data.get("confidence", 0.0))
        if not 0 <= confidence <= 1:
            raise TemplateFieldError("字段置信度必须在 0 到 1 之间")
        locator = dict(data.get("locator") or {})
        if source_format == "pdf":
            if locator.get("kind") != "pdf_rect":
                raise TemplateFieldError("PDF 字段必须提供坐标区域")
            rect = [float(value) for value in locator.get("rect", [])]
            if len(rect) != 4 or rect[2] <= rect[0] or rect[3] <= rect[1]:
                raise TemplateFieldError("PDF 字段坐标无效")
            locator["rect"] = rect
            locator["page"] = int(locator.get("page", -1))
            if locator["page"] < 0:
                raise TemplateFieldError("PDF 字段页码无效")
        normalized.append(
            TemplateField(
                key=key,
                label=str(data.get("label") or _label_for_key(key)).strip(),
                field_type=field_type,
                required=bool(data.get("required", False)),
                repeating=bool(data.get("repeating", False)),
                confidence=confidence,
                locator=locator,
            )
        )

    if source_format == "pdf":
        for index, field in enumerate(normalized):
            for other in normalized[index + 1 :]:
                if (
                    field.locator["page"] == other.locator["page"]
                    and _rectangles_overlap(
                        field.locator["rect"], other.locator["rect"]
                    )
                ):
                    raise TemplateFieldError(
                        f"PDF 字段区域重叠: {field.key}, {other.key}"
                    )
    return normalized


def parse_template(
    file_record: dict, infer_fields: InferenceCallback | None = None
) -> TemplateParseResult:
    extension = Path(file_record["safe_name"]).suffix.lower()
    if extension == ".docx":
        return _parse_docx(file_record["storage_path"], infer_fields)
    if extension == ".pdf":
        return _parse_pdf(file_record["storage_path"], infer_fields)
    raise TemplateFieldError(f"不支持的模板格式: {extension}")
