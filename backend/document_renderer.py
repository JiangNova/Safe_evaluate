"""Deterministic DOCX/PDF rendering and artifact packaging."""

from __future__ import annotations

import io
import json
import os
import shutil
import subprocess
import tempfile
import textwrap
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document

from .config import LIBREOFFICE_COMMAND
from .template_parser import PLACEHOLDER_RE, TemplateField


@dataclass(frozen=True)
class RenderWarning:
    code: str
    message: str
    field_key: str | None = None


@dataclass(frozen=True)
class RenderResult:
    path: str
    warnings: list[RenderWarning]


class DocumentRenderError(RuntimeError):
    """A document could not be rendered or converted."""


def render_compiled_template_docx(template_path, compiled, values, output_path, *, draft=True):
    """Lazy compatibility entry point for the universal DOCX renderer."""
    from .docx_renderer import render_compiled_docx

    return render_compiled_docx(template_path, compiled, values, output_path, draft=draft)


def _value_payload(values: dict, key: str) -> Any:
    raw = values.get(key, "")
    if hasattr(raw, "value"):
        return raw.value
    if isinstance(raw, dict) and "value" in raw:
        return raw["value"]
    return raw


def _display_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, list):
        return "\n".join(_display_value(item) for item in value)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _replace_paragraph(paragraph, replacements: dict[str, str]) -> set[str]:
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text or not PLACEHOLDER_RE.search(full_text):
        return set()
    used: set[str] = set()

    def replace(match):
        key = match.group(1)
        used.add(key)
        return replacements.get(key, "")

    rendered = PLACEHOLDER_RE.sub(replace, full_text)
    if paragraph.runs:
        paragraph.runs[0].text = rendered
        for run in paragraph.runs[1:]:
            run.text = ""
    else:  # pragma: no cover - python-docx normally creates a run for text
        paragraph.add_run(rendered)
    return used


def _replace_table(table, replacements: dict[str, str]) -> set[str]:
    used: set[str] = set()
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                used.update(_replace_paragraph(paragraph, replacements))
            for nested in cell.tables:
                used.update(_replace_table(nested, replacements))
    return used


def _fill_inferred_paragraph(paragraph, anchor: str, value: str) -> bool:
    """Insert an inferred value immediately after its confirmed anchor text."""
    full_text = "".join(run.text for run in paragraph.runs)
    if not anchor or anchor not in full_text:
        return False
    rendered = full_text.replace(anchor, f"{anchor}{value}", 1)
    if paragraph.runs:
        paragraph.runs[0].text = rendered
        for run in paragraph.runs[1:]:
            run.text = ""
    else:  # pragma: no cover
        paragraph.add_run(rendered)
    return True


def _fill_inferred_table(table, anchor: str, value: str) -> bool:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if _fill_inferred_paragraph(paragraph, anchor, value):
                    return True
            for nested in cell.tables:
                if _fill_inferred_table(nested, anchor, value):
                    return True
    return False


def render_docx(
    template_path: str,
    fields: list[TemplateField],
    values: dict,
    output_path: str,
) -> RenderResult:
    """Fill DOCX placeholders while retaining the first run's formatting."""
    document = Document(template_path)
    replacements = {
        field.key: _display_value(_value_payload(values, field.key))
        for field in fields
    }
    warnings: list[RenderWarning] = []
    for field in fields:
        rendered = replacements[field.key]
        if field.required and not rendered.strip():
            warnings.append(
                RenderWarning(
                    "missing_required_field",
                    f"必填字段“{field.label}”为空",
                    field.key,
                )
            )
        soft_limit = 2000 if field.field_type in {"multiline", "list"} else 500
        if len(rendered) > soft_limit:
            warnings.append(
                RenderWarning(
                    "field_overflow",
                    f"字段“{field.label}”内容较长，可能改变分页或表格高度",
                    field.key,
                )
            )

    used: set[str] = set()
    for paragraph in document.paragraphs:
        used.update(_replace_paragraph(paragraph, replacements))
    for table in document.tables:
        used.update(_replace_table(table, replacements))
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                used.update(_replace_paragraph(paragraph, replacements))
            for table in part.tables:
                used.update(_replace_table(table, replacements))

    for field in fields:
        if field.locator.get("kind") != "docx_inferred":
            continue
        anchor = str(field.locator.get("anchor", ""))
        value = replacements[field.key]
        found = any(
            _fill_inferred_paragraph(paragraph, anchor, value)
            for paragraph in document.paragraphs
        )
        if not found:
            found = any(
                _fill_inferred_table(table, anchor, value)
                for table in document.tables
            )
        if not found:
            for section in document.sections:
                for part in (section.header, section.footer):
                    if any(
                        _fill_inferred_paragraph(paragraph, anchor, value)
                        for paragraph in part.paragraphs
                    ) or any(
                        _fill_inferred_table(table, anchor, value)
                        for table in part.tables
                    ):
                        found = True
                        break
                if found:
                    break
        if not found:
            warnings.append(
                RenderWarning(
                    "anchor_not_found",
                    f"The confirmed anchor for field '{field.label}' was not found.",
                    field.key,
                )
            )

    for field in fields:
        if field.locator.get("kind") == "docx_placeholder" and field.key not in used:
            warnings.append(
                RenderWarning(
                    "placeholder_not_found",
                    f"模板中未找到字段“{field.label}”的占位符",
                    field.key,
                )
            )

    output = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output), exist_ok=True)
    document.save(output)
    return RenderResult(path=output, warnings=warnings)


def _wrap_pdf_text(text: str, rect: list[float], font_size: float) -> list[str]:
    width = max(1.0, rect[2] - rect[0])
    max_chars = max(1, int(width / max(font_size * 0.55, 1)))
    lines: list[str] = []
    for source_line in text.splitlines() or [""]:
        lines.extend(
            textwrap.wrap(
                source_line,
                width=max_chars,
                break_long_words=True,
                replace_whitespace=False,
            )
            or [""]
        )
    return lines


def render_pdf(
    template_path: str,
    fields: list[TemplateField],
    values: dict,
    output_path: str,
) -> RenderResult:
    """Overlay confirmed field rectangles onto the original PDF pages."""
    try:
        from pypdf import PdfReader, PdfWriter
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise DocumentRenderError("缺少 pypdf/reportlab，无法生成 PDF") from exc

    reader = PdfReader(template_path)
    if reader.is_encrypted:
        raise DocumentRenderError("暂不支持加密 PDF 模板")
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    overlay_stream = io.BytesIO()
    overlay = canvas.Canvas(overlay_stream)
    warnings: list[RenderWarning] = []

    fields_by_page: dict[int, list] = {}
    native_values: dict[str, Any] = {}
    for field in fields:
        if hasattr(field, "placements"):
            placement = next((item for item in field.placements if item.page is not None), None)
            if placement is None:
                continue
            page_number = int(placement.page)
            if placement.kind.startswith("pdf_form_"):
                raw = _value_payload(values, field.key)
                native_values[placement.anchor or field.key] = (
                    "/Yes" if placement.kind == "pdf_form_checkbox" and bool(raw)
                    else "/Off" if placement.kind == "pdf_form_checkbox"
                    else _display_value(raw)
                )
                continue
        else:
            page_number = int(field.locator.get("page", -1))
        if page_number < 0 or page_number >= len(reader.pages):
            raise DocumentRenderError(f"字段 {field.key} 的 PDF 页码无效")
        fields_by_page.setdefault(page_number, []).append(field)

    for page_index, page in enumerate(reader.pages):
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        overlay.setPageSize((width, height))
        for field in fields_by_page.get(page_index, []):
            if hasattr(field, "placements"):
                placement = next(item for item in field.placements if item.page == page_index)
                rect = [float(item) for item in (placement.rect or [])]
                font_size = float(placement.font_size)
                confirmed = placement.confirmed
            else:
                rect = [float(value) for value in field.locator["rect"]]
                font_size = float(field.locator.get("font_size", 10))
                confirmed = True
            if not confirmed:
                warnings.append(
                    RenderWarning(
                        "unconfirmed_pdf_placement",
                        f"字段“{field.label}”的 PDF 位置尚未确认",
                        field.key,
                    )
                )
                continue
            if (
                len(rect) != 4
                or rect[0] < 0
                or rect[1] < 0
                or rect[2] > width
                or rect[3] > height
            ):
                warnings.append(
                    RenderWarning(
                        "pdf_rect_out_of_bounds",
                        f"字段“{field.label}”超出 PDF 页面",
                        field.key,
                    )
                )
                continue
            value = _display_value(_value_payload(values, field.key))
            if field.required and not value.strip():
                warnings.append(
                    RenderWarning(
                        "missing_required_field",
                        f"必填字段“{field.label}”为空",
                        field.key,
                    )
                )
            line_height = font_size * 1.25
            lines = _wrap_pdf_text(value, rect, font_size)
            available_lines = max(1, int((rect[3] - rect[1]) / line_height))
            if len(lines) > available_lines:
                warnings.append(
                    RenderWarning(
                        "pdf_text_overflow" if hasattr(field, "placements") else "field_overflow",
                        f"字段“{field.label}”内容超出 PDF 填写区域",
                        field.key,
                    )
                )
            overlay.setFont("STSong-Light", font_size)
            y = rect[3] - font_size
            for line in lines[:available_lines]:
                overlay.drawString(rect[0], y, line)
                y -= line_height
        overlay.showPage()
    overlay.save()
    overlay_stream.seek(0)

    overlay_reader = PdfReader(overlay_stream)
    writer = PdfWriter(clone_from=template_path)
    for page_index, page in enumerate(writer.pages):
        page.merge_page(overlay_reader.pages[page_index])
    if native_values:
        for page in writer.pages:
            try:
                writer.update_page_form_field_values(
                    page, native_values, auto_regenerate=True
                )
            except Exception as exc:
                warnings.append(
                    RenderWarning(
                        "pdf_form_appearance_failed",
                        f"PDF 表单外观更新失败: {exc}",
                    )
                )
    output = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output), exist_ok=True)
    with open(output, "wb") as stream:
        writer.write(stream)
    return RenderResult(path=output, warnings=warnings)


def convert_docx_to_pdf(
    docx_path: str, output_dir: str, timeout: int = 120
) -> Path:
    """Convert DOCX with an isolated LibreOffice profile."""
    executable = shutil.which(LIBREOFFICE_COMMAND)
    if not executable:
        raise DocumentRenderError("服务器未安装 LibreOffice，无法转换 PDF")
    output = Path(output_dir).resolve()
    output.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="safe-evaluate-lo-") as profile:
        profile_uri = Path(profile).resolve().as_uri()
        command = [
            executable,
            "--headless",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output),
            str(Path(docx_path).resolve()),
        ]
        try:
            completed = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=timeout,
                check=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise DocumentRenderError("LibreOffice 转换超时") from exc
    expected = output / f"{Path(docx_path).stem}.pdf"
    if completed.returncode != 0 or not expected.exists():
        detail = (completed.stderr or completed.stdout or "unknown error").strip()
        raise DocumentRenderError(f"LibreOffice 转换失败: {detail[:500]}")
    return expected


def build_artifact_zip(
    files: Iterable[tuple[str, str]],
    failures: list[dict],
    output_path: str,
) -> str:
    """Package finalized artifacts and a readable partial-failure manifest."""
    output = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output), exist_ok=True)
    used_names: set[str] = set()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path, requested_name in files:
            name = os.path.basename(requested_name) or os.path.basename(path)
            if name in used_names:
                stem, suffix = os.path.splitext(name)
                counter = 2
                while f"{stem}-{counter}{suffix}" in used_names:
                    counter += 1
                name = f"{stem}-{counter}{suffix}"
            used_names.add(name)
            archive.write(path, name)
        if failures:
            lines = ["以下文书未能加入压缩包：", ""]
            for failure in failures:
                lines.append(
                    f"- {failure.get('template', '未知模板')}: {failure.get('error', '未知错误')}"
                )
            archive.writestr("失败清单.txt", "\n".join(lines).encode("utf-8"))
    return output
