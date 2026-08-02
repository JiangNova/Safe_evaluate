"""Styled DOCX rendering for compiled text-defined templates."""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from .document_renderer import RenderResult, RenderWarning
from .docx_renderer import MISSING_MARKER
from .template_ir import CompiledTemplate


def _value(values: dict, key: str) -> str:
    raw = values.get(key, "")
    if isinstance(raw, dict):
        raw = raw.get("value", "")
    if isinstance(raw, list):
        return "\n".join(str(item) for item in raw)
    return str(raw or "")


def render_text_document(
    compiled: CompiledTemplate,
    values: dict,
    output_docx: str,
    *,
    draft: bool = True,
) -> RenderResult:
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.left_margin = section.right_margin = Mm(25)
    section.top_margin, section.bottom_margin = Mm(24), Mm(22)
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    title = document.add_heading(compiled.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if draft:
        draft_mark = document.add_paragraph("【草稿 · 尚未完成质量校验】")
        draft_mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warnings = []
    for field in compiled.fields:
        document.add_heading(field.label, level=1)
        rendered = _value(values, field.key)
        if not rendered and field.required:
            warnings.append(RenderWarning("missing_required_field", f"必填字段“{field.label}”为空", field.key))
            rendered = MISSING_MARKER if draft else ""
        paragraph = document.add_paragraph(rendered)
        paragraph.paragraph_format.space_after = Pt(10)
        raw = values.get(field.key)
        refs = raw.get("source_refs", []) if isinstance(raw, dict) else []
        if refs:
            document.add_paragraph(f"来源：{'；'.join(map(str, refs))}", style="Caption")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("草稿 · 通用自动评估平台" if draft else "通用自动评估平台")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    document.save(output_docx)
    return RenderResult(output_docx, warnings)
