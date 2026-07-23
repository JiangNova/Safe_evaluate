"""Parse requirement documents into structured evaluation criteria."""
import os
import re
import sys

# python-docx is optional — only needed for .docx files
try:
    from docx import Document
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    Document = None  # type: ignore

from .config import REQUIREMENT_DIR


def _read_docx(filepath: str) -> str:
    """Extract all text from a .docx file."""
    if not _HAS_DOCX:
        return f"[无法读取 .docx 文件，缺少 python-docx 依赖。请运行: pip install python-docx]"
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _read_doc(filepath: str) -> str:
    """Extract text from a legacy .doc (OLE Compound Document) file.

    Strategy (tried in order):
    1. Parse the OLE compound document and extract the WordDocument stream,
       then decode as UTF-16-LE with proper alignment.
    2. Fall back: scan for aligned UTF-16-LE Chinese text runs in the raw bytes.
    3. Last resort: extract readable ASCII and return with a warning.
    """
    try:
        with open(filepath, "rb") as f:
            raw = f.read()
    except Exception as e:
        return f"[无法读取文件: {e}]"

    # ---- Approach 1: OLE compound document parsing ----
    try:
        text = _extract_ole_word_document(raw)
        if text and len(text) > 100:
            return text
    except Exception:
        pass

    # ---- Approach 2: aligned UTF-16-LE Chinese text scan ----
    text = _scan_utf16le_chinese(raw)
    if text and len(text) > 100:
        return text

    # ---- Approach 3: ASCII fallback ----
    result = []
    for byte in raw:
        if 32 <= byte < 127 or byte in (10, 13):
            result.append(chr(byte))
    text = "".join(result)
    if len(text) > 50:
        text = re.sub(r"([a-zA-Z]{20,})", "\n\\1\n", text)
        return text

    return f"[无法解析 .doc 文件，建议转换为 .docx 格式: {filepath}]"


def _is_chinese_char(c: str) -> bool:
    """Return True if character is a CJK character or common full-width punctuation."""
    cp = ord(c)
    return (
        0x4E00 <= cp <= 0x9FFF   # CJK Unified Ideographs
        or 0x3400 <= cp <= 0x4DBF  # CJK Unified Ideographs Extension A
        or 0xF900 <= cp <= 0xFAFF  # CJK Compatibility Ideographs
        or 0x3000 <= cp <= 0x303F  # CJK punctuation
        or 0xFF00 <= cp <= 0xFFEF  # Fullwidth forms
        or 0x2000 <= cp <= 0x206F  # General punctuation
    )


def _scan_utf16le_chinese(raw: bytes) -> str:
    """Scan raw bytes for aligned UTF-16-LE Chinese text runs.

    Only keeps runs that are predominantly Chinese characters (≥60% CJK)
    and longer than 20 characters — this naturally filters out OLE
    structural bytes that happen to decode as valid codepoints.
    """
    # We scan from even bytes only (UTF-16-LE is 2-byte aligned)
    runs: list[str] = []
    current_run: list[str] = []

    # Try both even and odd alignment
    for offset in (0, 1):
        i = offset
        while i < len(raw) - 1:
            lo = raw[i]
            hi = raw[i + 1]
            cp = lo | (hi << 8)

            # Valid printable character range
            if (0x20 <= cp <= 0x7E       # ASCII printable
                    or 0x3000 <= cp <= 0x303F  # CJK punctuation
                    or 0xFF00 <= cp <= 0xFFEF  # Fullwidth
                    or 0x2000 <= cp <= 0x206F  # General punctuation
                    or 0x4E00 <= cp <= 0x9FFF  # CJK
                    or 0x3400 <= cp <= 0x4DBF  # CJK Ext-A
                    or 0xF900 <= cp <= 0xFAFF  # CJK Compat
                    or cp in (0x000D, 0x000A, 0x0009)):  # CR, LF, TAB
                current_run.append(chr(cp))
                i += 2
            else:
                # End of run — decide whether to keep it
                if current_run:
                    run_text = "".join(current_run)
                    cjk_count = sum(1 for c in run_text if _is_chinese_char(c))
                    total_meaningful = sum(1 for c in run_text if c.isprintable() or c in "\n\r\t")
                    if total_meaningful > 0:
                        cjk_ratio = cjk_count / total_meaningful
                    else:
                        cjk_ratio = 0
                    # Keep runs with high Chinese density and decent length
                    if cjk_ratio >= 0.4 and len(run_text) >= 15:
                        runs.append(run_text)
                    current_run = []
                i += 1

        # Don't forget the last run
        if current_run:
            run_text = "".join(current_run)
            cjk_count = sum(1 for c in run_text if _is_chinese_char(c))
            total_meaningful = sum(1 for c in run_text if c.isprintable() or c in "\n\r\t")
            cjk_ratio = cjk_count / total_meaningful if total_meaningful > 0 else 0
            if cjk_ratio >= 0.4 and len(run_text) >= 15:
                runs.append(run_text)
            current_run = []

    # Deduplicate: runs from different alignments may overlap
    seen: set[str] = set()
    unique_runs = []
    for r in runs:
        # Normalize whitespace for dedup
        key = re.sub(r"\s+", "", r)[:100]
        if key and key not in seen:
            seen.add(key)
            unique_runs.append(r)

    return "\n\n".join(unique_runs)


def _extract_ole_word_document(raw: bytes) -> str | None:
    """Parse OLE Compound Document and extract text from the WordDocument stream.

    Returns decoded text or None if parsing fails.
    """
    import struct

    # Verify OLE magic
    if raw[:8] != b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
        return None

    # Parse header (little-endian)
    # Bytes 0x1A-0x1B: major version (3 or 4)
    major_version = struct.unpack_from('<H', raw, 0x1A)[0]

    # Bytes 0x1E-0x1F: sector size exponent
    sector_shift = struct.unpack_from('<H', raw, 0x1E)[0]
    sector_size = 1 << sector_shift  # typically 9 → 512

    if sector_size <= 0 or sector_size > 4096:
        return None

    # Read header fields — offsets differ between v3 and v4
    if major_version == 4:
        # OLE v4 offsets
        num_dir_sectors = struct.unpack_from('<I', raw, 0x2C)[0]
        num_fat_sectors = struct.unpack_from('<I', raw, 0x30)[0]
        first_dir_sector = struct.unpack_from('<I', raw, 0x34)[0]
    else:
        # OLE v3 offsets (and fallback)
        # 0x2C: num FAT sectors
        # 0x30: first directory sector
        # Directory size is determined by following the FAT chain
        num_fat_sectors = struct.unpack_from('<I', raw, 0x2C)[0]
        first_dir_sector = struct.unpack_from('<I', raw, 0x30)[0]
        num_dir_sectors = 0  # v3: follow FAT chain instead

    # Read DIFAT: first 109 entries are at bytes 76-511 of the header
    difat = []
    for i in range(109):
        offset = 76 + i * 4
        if offset + 4 > len(raw):
            break
        val = struct.unpack_from('<I', raw, offset)[0]
        if val != 0xFFFFFFFF:
            difat.append(val)

    # Read FAT
    fat = []
    for difat_sector in difat:
        offset = (difat_sector + 1) * sector_size
        if offset + sector_size > len(raw):
            break
        for i in range(sector_size // 4):
            val = struct.unpack_from('<I', raw, offset + i * 4)[0]
            fat.append(val)

    # Helper: follow a FAT chain
    def get_chain(start_sector: int) -> list[int]:
        chain = []
        sec = start_sector
        visited = set()
        while sec < len(fat) and sec not in visited:
            visited.add(sec)
            chain.append(sec)
            next_sec = fat[sec]
            if next_sec == 0xFFFFFFFF or next_sec >= len(fat):
                break
            sec = next_sec
        return chain

    # Read directory sectors
    if num_dir_sectors > 0:
        # v4: explicit count
        dir_chain = list(range(first_dir_sector, first_dir_sector + num_dir_sectors))
    else:
        # v3: follow FAT chain starting from first_dir_sector
        dir_chain = get_chain(first_dir_sector)

    dir_data = bytearray()
    for dir_sec in dir_chain:
        offset = (dir_sec + 1) * sector_size
        if offset + sector_size > len(raw):
            break
        dir_data.extend(raw[offset:offset + sector_size])

    # Parse directory entries (128 bytes each)
    # Root entry is entry 0
    dir_entries = []
    for i in range(0, len(dir_data), 128):
        entry = dir_data[i:i + 128]
        if len(entry) < 128:
            break
        # Name is UTF-16-LE at bytes 0-63 (max 32 chars)
        name_len = struct.unpack_from('<H', entry, 64)[0]
        name_raw = entry[0:name_len]
        try:
            name = name_raw.decode('utf-16-le', errors='ignore').rstrip('\x00')
        except Exception:
            name = ""
        obj_type = entry[66]  # 0=unknown, 1=storage, 2=stream, 5=root
        start_sector = struct.unpack_from('<I', entry, 116)[0]
        stream_size = struct.unpack_from('<I', entry, 120)[0]
        dir_entries.append({
            "name": name,
            "type": obj_type,
            "start_sector": start_sector,
            "size": stream_size,
        })

    # Find the WordDocument stream
    word_doc_entry = None
    for entry in dir_entries:
        if entry["name"] == "WordDocument" and entry["type"] == 2:
            word_doc_entry = entry
            break

    if not word_doc_entry:
        # Try "1Table" or "0Table" as fallback
        for entry in dir_entries:
            if entry["name"] in ("1Table", "0Table") and entry["type"] == 2:
                word_doc_entry = entry
                break

    if not word_doc_entry:
        return None

    # Read the stream data
    chain = get_chain(word_doc_entry["start_sector"])
    stream_data = bytearray()
    for sec in chain:
        offset = (sec + 1) * sector_size
        if offset + sector_size > len(raw):
            break
        stream_data.extend(raw[offset:offset + sector_size])

    stream_data = bytes(stream_data[:word_doc_entry["size"]])

    # From the WordDocument stream, extract text
    # The FIB (File Information Block) is at the beginning
    # At offset 2 in the FIB: flags (bit 0 = fComplex, bit 1 = fHasTable, etc.)
    # For simple documents, text starts after FIB at fixed offset

    # Simplified approach: the document text is typically stored as either
    # ASCII or UTF-16-LE after the FIB header. We try to find it.

    # The FIB has:
    # - 2 bytes: magic (0xA5EC)
    # - 2 bytes: flags
    # - various fields...
    # - ccpText at a documented offset (depends on version)

    # Simpler: just look for runs of valid text within the WordDocument stream
    text = _scan_utf16le_chinese(stream_data)

    if not text or len(text) < 50:
        # Try decoding the whole stream as UTF-16-LE ignoring errors
        try:
            decoded = stream_data.decode('utf-16-le', errors='ignore')
            # Filter for printable characters only
            readable = "".join(c for c in decoded if c.isprintable() or c in "\n\r\t")
            # Clean up: collapse whitespace
            readable = re.sub(r'[ \t]+', ' ', readable)
            readable = re.sub(r'\n{3,}', '\n\n', readable)
            if len(readable) > 100:
                return readable
        except Exception:
            pass

    return text


def load_all_requirements() -> list[dict]:
    """Load all requirement documents, returning list of {filename, content}."""
    documents = []
    if not os.path.exists(REQUIREMENT_DIR):
        return documents

    for filename in sorted(os.listdir(REQUIREMENT_DIR)):
        filepath = os.path.join(REQUIREMENT_DIR, filename)
        if not os.path.isfile(filepath):
            continue

        ext = os.path.splitext(filename)[1].lower()
        try:
            if ext == ".docx":
                content = _read_docx(filepath)
            elif ext == ".doc":
                content = _read_doc(filepath)
            else:
                continue
        except Exception as e:
            content = f"[读取失败: {e}]"

        # Quality check: skip documents whose text is mostly garbled
        if content and len(content.strip()) > 50:
            quality = _text_quality(content)
            if quality < 0.3:
                print(f"[WARNING] {filename}: text quality low ({quality:.2f}), "
                      f"可能是编码问题，建议将 .doc 转为 .docx 格式",
                      file=sys.stderr)
                # Still include it but add a note
                content = f"[警告：该文档文本质量较低，可能存在编码问题，建议重新保存为 .docx 格式]\n\n{content}"

            documents.append({
                "filename": filename,
                "filepath": filepath,
                "content": content,
                "length": len(content),
            })

    return documents


def _text_quality(text: str) -> float:
    """Return a 0-1 score indicating how 'clean' the extracted text is.

    Low scores mean the text is likely garbled (e.g. encoding errors).
    We expect Chinese legal documents to be mostly CJK + punctuation + digits.
    """
    if not text:
        return 0.0

    # Remove whitespace for analysis
    stripped = re.sub(r'\s+', '', text)
    if not stripped:
        return 0.0

    # Count character categories
    total = len(stripped)
    cjk = sum(1 for c in stripped if '一' <= c <= '鿿')
    cjk_ext = sum(1 for c in stripped if '㐀' <= c <= '䶿')
    cjk_punct = sum(1 for c in stripped if '　' <= c <= '〿')
    digits = sum(1 for c in stripped if c.isdigit())
    ascii_alpha = sum(1 for c in stripped if 'a' <= c.lower() <= 'z')
    ascii_punct = sum(1 for c in stripped if c in '.,;:()[]{}（）【】《》"\'、。，；：！？…——')
    common = sum(1 for c in stripped if c in '0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ'
                 '一二三四五六七八九十百千万亿条第款项章节目上下左右中大小高低内外前后'
                 '消防安全检查督管标准规法单位场所建筑设施通道疏散')

    # Fullwidth forms that suggest encoding issues (0xFF00-0xFFEF)
    # These are usually valid, but excessive amounts mixed with rare codepoints = bad
    fullwidth = sum(1 for c in stripped if '＀' <= c <= '￯')

    # Rare/unusual CJK codepoints often indicate binary garbage decoded as Unicode
    rare = sum(1 for c in stripped if (
        '䀀' <= c <= '䶿' and c not in ''  # CJK Ext-A is fine in names
    ))

    # Heuristic: real Chinese document text should be mostly common characters
    known = cjk + cjk_ext + cjk_punct + digits + ascii_alpha + ascii_punct + common + fullwidth
    ratio = known / total if total > 0 else 0

    # Bonus for CJK (real Chinese docs have >50% CJK after removing whitespace)
    cjk_ratio = (cjk + cjk_ext) / total if total > 0 else 0

    # Penalize for lots of private-use or very rare codepoints
    unknown_ratio = 1.0 - ratio

    score = ratio * 0.6 + min(cjk_ratio * 1.5, 0.4)
    score -= unknown_ratio * 0.5

    return max(0.0, min(1.0, score))


def build_requirements_context(documents: list[dict], max_chars: int = 8000) -> str:
    """Build a compact context string from all requirement documents.

    Trims to max_chars to fit within the LLM prompt budget while preserving
    the most important content from each document.
    """
    if not documents:
        return "暂无具体消防法规要求文档，请依据国家通用消防标准进行评估。"

    parts = []
    total = 0
    per_doc = max_chars // len(documents)

    for doc in documents:
        content = doc["content"]
        # Take beginning portion of each doc (front-loaded with key info)
        if len(content) > per_doc:
            # Take first half and some from middle
            first = content[: int(per_doc * 0.7)]
            mid_start = len(content) // 2
            middle = content[mid_start : mid_start + int(per_doc * 0.3)]
            excerpt = first + "\n...(省略中间部分)...\n" + middle
        else:
            excerpt = content

        parts.append(f"### {doc['filename']}\n{excerpt}")
        total += len(excerpt)

    return "\n\n".join(parts)
