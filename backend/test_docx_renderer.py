import os
import tempfile
import unittest

from docx import Document

from backend.docx_renderer import render_compiled_docx
from backend.docx_template_compiler import compile_docx_template


class DocxRendererTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()

    def tearDown(self):
        self.temp.cleanup()

    def test_replaces_underlined_blank_in_place(self):
        template = os.path.join(self.temp.name, "template.docx")
        output = os.path.join(self.temp.name, "output.docx")
        document = Document()
        paragraph = document.add_paragraph("员工：")
        blank = paragraph.add_run("________")
        blank.font.underline = True
        paragraph.add_run(" 因违纪")
        document.save(template)
        compiled = compile_docx_template(template)
        field = compiled.fields[0]
        render_compiled_docx(template, compiled, {field.key: {"value": "张三"}}, output)
        self.assertEqual(Document(output).paragraphs[0].text, "员工：张三 因违纪")

    def test_selects_exact_checkbox_option(self):
        template = os.path.join(self.temp.name, "checkbox.docx")
        output = os.path.join(self.temp.name, "output.docx")
        document = Document()
        document.add_paragraph("处罚：□警告  □记过  □辞退")
        document.save(template)
        compiled = compile_docx_template(template)
        field = compiled.fields[0]
        render_compiled_docx(template, compiled, {field.key: {"value": "记过"}}, output)
        self.assertIn("□警告  ☑记过  □辞退", Document(output).paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()
